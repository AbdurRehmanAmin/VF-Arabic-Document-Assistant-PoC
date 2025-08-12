# document_processor.py
import PyPDF2
import docx
import re
from typing import List, Dict, Any, Union
import os
from langdetect import detect
import logging
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

async def process_document(file_path: str) -> str:
    """Extract text from documents in various formats with improved error handling."""
    logger.info(f"Processing document: {file_path}")
    
    try:
        if file_path.endswith('.pdf'):
            return extract_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return extract_from_docx(file_path)
        elif file_path.endswith('.txt'):
            return extract_from_txt(file_path)
        else:
            return "صيغة الملف غير مدعومة"
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return f"خطأ في معالجة المستند: {str(e)}"

def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF with special handling for Arabic text."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Log number of pages for debugging
            logger.info(f"PDF has {len(reader.pages)} pages")
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                logger.info(f"Page {i+1} extracted {len(page_text) if page_text else 0} characters")
                if page_text:
                    text += page_text + "\n\n"
    except Exception as e:
        logger.error(f"Error extracting PDF: {str(e)}")
        raise
    
    # Handle right-to-left text and fix common PDF extraction issues with Arabic
    text = fix_arabic_text(text)
    
    # Debug log the extracted text length
    logger.info(f"Extracted {len(text)} characters from PDF")
    return text

def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX with improved extraction."""
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Extract paragraphs with better formatting
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                text += para.text.strip() + "\n\n"
        
        # Also extract tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
            text += "\n"
    except Exception as e:
        logger.error(f"Error extracting DOCX: {str(e)}")
        raise
    
    # Handle right-to-left text
    text = fix_arabic_text(text)
    
    # Debug log the extracted text length
    logger.info(f"Extracted {len(text)} characters from DOCX")
    return text

def extract_from_txt(file_path: str) -> str:
    """Extract text from TXT with encoding detection."""
    # Try different encodings, starting with UTF-8
    encodings = ['utf-8', 'utf-16', 'cp1256', 'iso-8859-6', 'windows-1256']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
                logger.info(f"Successfully read text file with encoding: {encoding}")
                # If we successfully read the file, fix Arabic text
                return fix_arabic_text(text)
        except UnicodeDecodeError:
            continue
    
    # If all encodings fail, try binary mode and then decode
    with open(file_path, 'rb') as file:
        binary_data = file.read()
        # Try to detect encoding from binary data
        for encoding in encodings:
            try:
                text = binary_data.decode(encoding)
                logger.info(f"Successfully decoded binary data with encoding: {encoding}")
                return fix_arabic_text(text)
            except UnicodeDecodeError:
                continue
    
    logger.error("Failed to decode text file with any standard encoding")
    return "لم يتمكن من قراءة الملف النصي باستخدام أي ترميز قياسي."

def fix_arabic_text(text: str) -> str:
    """Fix common issues with Arabic text extraction."""
    if not text or len(text.strip()) == 0:
        return ""
        
    # Remove extra whitespace but preserve paragraph breaks
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Fix common Arabic ligature issues
    text = text.replace('ﻻ', 'لا')
    
    # Normalize Arabic characters (alef variations, etc.)
    replacements = {
        'أ': 'ا',
        'إ': 'ا',
        'آ': 'ا',
        'ى': 'ي',
        'ة': 'ه'
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    return text

def detect_language(text: str) -> str:
    """Detect if text is primarily Arabic or English."""
    if not text or len(text.strip()) == 0:
        return 'unknown'
        
    # First check for Arabic characters
    arabic_char_count = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
    if arabic_char_count > len(text) * 0.3:  # If more than 30% is Arabic
        return 'arabic'
    
    # Use langdetect as a fallback
    try:
        lang = detect(text[:1000])  # Use first 1000 chars for faster detection
        if lang == 'ar':
            return 'arabic'
        else:
            return 'english'
    except:
        # Default to Arabic if detection fails (since target users are Arabic speakers)
        return 'arabic'

async def chunk_document(text: str) -> List[Dict[str, Any]]:
    """
    Split document into semantically meaningful chunks using LangChain's text splitters.
    Returns chunks with metadata for better retrieval including page info.
    """
    logger.info("Using semantic chunking with LangChain")
    
    # Check if text is mostly Arabic
    is_arabic = any('\u0600' <= c <= '\u06FF' for c in text[:500])
    
    if not text or len(text.strip()) == 0:
        logger.warning("Empty document text provided for chunking")
        return []
    
    try:
        # Split the text by page markers (assuming each page break is marked by form feeds or specific patterns)
        pages = []
        
        # Try to split by form feed characters first
        if '\f' in text:
            pages = text.split('\f'),
        else:
            # If no form feeds, try to infer pages by double newlines or specific patterns
            pages = re.split(r'\n\s*\n\s*\n', text)
        
        # If the split resulted in too few or too many pages, fall back to a simpler approach
        if len(pages) < 2 or len(pages) > 500:
            pages = [text]
        
        # Use RecursiveCharacterTextSplitter for better semantic chunking
        if is_arabic:
            # For Arabic, use appropriate separators
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,  # Smaller chunks for better context
                chunk_overlap=100,
                separators=["\n\n", "\n", ".", "!", "?", "،", "؛", " ", ""],
            )
        else:
            # For English and other languages
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=100,
                separators=["\n\n", "\n", ".", "!", "?", ";", ":", " ", ""],
            )
        
        # Create chunks with page information
        chunks = []
        chunk_id = 0
        
        for page_num, page_text in enumerate(pages):
            if not page_text.strip():
                continue
                
            # Create Document object with text and metadata
            doc = Document(page_content=page_text, metadata={"page": page_num + 1})
            
            # Split the page into chunks
            page_chunks = text_splitter.split_documents([doc])
            
            # Extract line numbers for each chunk
            for i, chunk in enumerate(page_chunks):
                # Estimate line number based on position in page
                # This is approximate since we don't have exact line breaks
                start_pos = page_text.find(chunk.page_content[:50])
                
                if start_pos == -1:
                    line_num = 1  # Default if can't determine
                else:
                    # Count newlines before this chunk to estimate line number
                    line_num = page_text[:start_pos].count('\n') + 1
                
                chunks.append({
                    "id": chunk_id,
                    "text": chunk.page_content,
                    "metadata": {
                        "source": "document",
                        "chunk_id": chunk_id,
                        "page": page_num + 1,
                        "line": line_num,
                        "language": "arabic" if is_arabic else "english"
                    }
                })
                chunk_id += 1
        
        logger.info(f"Document split into {len(chunks)} semantic chunks across {len(pages)} pages")
        return chunks
    except Exception as e:
        logger.error(f"Error during document chunking: {str(e)}")
        # Return single chunk as fallback
        return [{"id": 0, "text": text, "metadata": {"source": "document", "chunk_id": 0, "page": 1, "line": 1}}]